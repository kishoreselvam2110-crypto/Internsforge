import io
import pdfplumber
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from PDF bytes using pdfplumber.
    """
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_text = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
            text = "\n".join(pages_text)
    except Exception as e:
        print(f"pdfplumber extraction failed: {str(e)}")
        text = ""

    text = text.strip()
    if not text:
        raise ValueError("Could not extract any text from the PDF. The file might be scanned, corrupt, or empty.")
    
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX bytes using python-docx.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        
        # Also extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        tables_text.append(cell.text.strip())
                        
        combined_text = "\n".join(paragraphs + tables_text)
        combined_text = combined_text.strip()
        if not combined_text:
            raise ValueError("DOCX file contains no text.")
        return combined_text
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Helper function to dispatch extraction based on file extension.
    """
    lower_filename = filename.lower()
    if lower_filename.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_filename.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type for filename '{filename}'. Only PDF and DOCX files are supported.")

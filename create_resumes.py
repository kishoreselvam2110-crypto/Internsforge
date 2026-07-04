import os
try:
    from docx import Document
except ImportError:
    print("python-docx not installed yet.")
    exit(1)

os.makedirs("test_resumes", exist_ok=True)

# 1. The Perfect Match
doc1 = Document()
doc1.add_heading('Alice Smith', 0)
doc1.add_paragraph('alice.smith@email.com | +1 555-0100')
doc1.add_heading('Summary', level=1)
doc1.add_paragraph('Senior Software Engineer with 6 years of professional experience building scalable web applications. Highly proficient in Python, FastAPI, React, and TypeScript.')
doc1.add_heading('Experience', level=1)
doc1.add_paragraph('Lead Developer at TechCorp')
doc1.add_paragraph('January 2020 to Present')
doc1.add_paragraph('Developed high-performance microservices using FastAPI and PostgreSQL. Designed robust frontends using React and Next.js. Deployed on AWS using Docker and Kubernetes.')
doc1.add_heading('Education', level=1)
doc1.add_paragraph("Master's in Computer Science, University of Technology")
doc1.save('test_resumes/alice_smith_perfect_match.docx')

# 2. The Partial Match
doc2 = Document()
doc2.add_heading('Bob Jones', 0)
doc2.add_paragraph('bob.j@email.com | (555) 0199')
doc2.add_heading('Summary', level=1)
doc2.add_paragraph('Backend Developer with 3 years of experience. Passionate about learning new technologies.')
doc2.add_heading('Experience', level=1)
doc2.add_paragraph('Junior Developer at StartupInc')
doc2.add_paragraph('March 2023 - Present')
doc2.add_paragraph('Maintained backend systems written in Django and Python. Wrote SQL queries and managed MongoDB instances. Familiar with Git and Agile methodologies.')
doc2.add_heading('Education', level=1)
doc2.add_paragraph("Bachelor's Degree in Software Engineering, State College")
doc2.save('test_resumes/bob_jones_partial_match.docx')

# 3. The Poor Match
doc3 = Document()
doc3.add_heading('Charlie Brown', 0)
doc3.add_paragraph('charlie.b@email.com | 555-888-0000')
doc3.add_heading('Summary', level=1)
doc3.add_paragraph('Recent graduate looking for an entry-level marketing position. Experienced with social media campaigns and content creation.')
doc3.add_heading('Experience', level=1)
doc3.add_paragraph('Social Media Intern')
doc3.add_paragraph('June 2025 to August 2025 (0.2 years)')
doc3.add_paragraph('Created graphics using Canva and Figma. Managed Instagram and Twitter accounts. Wrote SEO optimized blog posts.')
doc3.add_heading('Education', level=1)
doc3.add_paragraph("High School Diploma, Springfield High")
doc3.save('test_resumes/charlie_brown_poor_match.docx')

print("Successfully created 3 test resumes in the 'test_resumes' directory.")
